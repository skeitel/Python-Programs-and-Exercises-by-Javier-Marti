'''This program extracts information from different websites to create a unique report in the form of a Word document'''

from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches

#Get paragraph 1
url = requests.get('https://en.wikipedia.org/wiki/Formula_1')
soup = BeautifulSoup(url.content, 'lxml')
paragraph1 = soup.find_all('p')[1].text
print(paragraph1)

#Get paragraph 2
url = requests.get('https://auto.howstuffworks.com/auto-racing/motorsports/formula-one.htm')
soup = BeautifulSoup(url.content, 'lxml')
paragraph2 = soup.find_all('p')[5].text[1:]
print(paragraph2)

#Get paragraph 3
url = requests.get('https://www.statesman.com/news/local/what-formula-one/9ptnv9xDhnVExEkYLGBivL/')
soup = BeautifulSoup(url.content, 'lxml')
paragraph3 = soup.find_all('p')[8].text
print('\n', paragraph3)

#Create Word doc
document = Document()
document.add_heading('Formula 1', 0)

#add picture
document.add_picture('f1car.jpg', width= Inches(5))

#add paragraphs
p0 = document.add_paragraph()
p1 = document.add_paragraph()
p1.add_run(paragraph1).bold = True
p2= document.add_paragraph(paragraph2)
p3= document.add_paragraph(paragraph3)

#save doc
filename = 'f1report.docx'
document.save(filename)

print('.'*80)
print('\nThe document has created and saved as {}'.format(filename))